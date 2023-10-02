from tkinter import *
from tkinter.filedialog import askopenfile,asksaveasfile
import docx
from openpyxl import load_workbook
import io
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.asymmetric import dsa
from cryptography.hazmat.primitives import serialization
from DSA import Kyso, Xacnhankyso

FontBig = ('Time New Roman',15,'bold')
FontSmall = ('Time New Roman',13,'bold')
BG = '#FFFAFA'

window = Tk()
window.title("Ninh Duy Nhật - Ký số DSA")
window.geometry('1600x700')
#Mã hóa
label0 = Label(window, text='Phát sinh chữ ký', font=FontBig, bg=BG)
label0.place(x=1, y=1)

label1 = Label(window, text='Chọn L:', font=FontSmall, bg=BG)
label1.place(x=1,y=35)

label2 = Label(window, text='Văn bản ký:', font=FontSmall, bg=BG)
label2.place(x=1,y=65)

label3 = Label(window, text='Chữ ký:', font=FontSmall, bg=BG)
label3.place(x=1,y=300)

key_entry = Entry(window, width=20, font=FontSmall)
key_entry.place(x=100, y=35)

p_entry = Text(window, width=50, height=10)
p_entry.place(x=1, y=90)

ans1_entry = Text(window, width=50, height=10)
ans1_entry.place(x=1, y=325)

#Giải mã
label0 = Label(window, text='Kiểm tra chữ ký', font=FontBig, bg=BG)
label0.place(x=700, y=1)

label6 = Label(window, text='Văn bản ký:', font=FontSmall, bg=BG)
label6.place(x=700,y=65)

label7 = Label(window, text='Chữ ký', font=FontSmall, bg=BG)
label7.place(x=700,y=300)

label8 = Label(window, text='Thông báo', font=FontSmall, bg=BG)
label8.place(x=700,y=500)

ans2_entry = Text(window, width=50, height=10)
ans2_entry.place(x=700, y=90)

p1_entry = Text(window, width=50, height=10)
p1_entry.place(x=700, y=325)

p2_entry = Text(window, width=50, height=2)
p2_entry.place(x=700, y=550)

#Mở các file đề thi
def readWordfile(filename):
    doc = docx.Document(filename)
    result = [p.text for p in doc.paragraphs]
    return '\n'.join(result)


def readExcelfile(filename):
    book = load_workbook(filename)
    sheet = book.active
    ans = ''
    for i in range(1,sheet.max_row+1):
        for j in range(1,sheet.max_column+1):
            ans+= str(sheet.cell(row=i, column=j).value)
            ans+=' '
        ans+='\n'
    return ans

def readTextfile(filename):
    s= ''
    file = io.open(filename, mode='r', encoding='utf-8')
    read = file.readlines()
    for line in read:
        s+=line
    return s

filetypes = [('Word Files', '*.docx'), ('Text Files', '*.txt'), ('Excel Files', '*.xlsx')]
def open():
    global data
    file = askopenfile(mode='r', initialdir='C:/Users/Admin/Desktop/An Toan Va Bao Mat Thong Tin/BTL', title='Select File', filetypes=filetypes)
    s = str(file)
    a = s.find('name')
    b = s.find('mode')
    filename = s[a+6:b-2]
    if 'docx' in filename:
        data = readWordfile(filename)
    elif 'xlsx' in filename:
        data = readExcelfile(filename)
    elif 'txt' in filename:
        data = readTextfile(filename)
    p_entry.delete(1.0,END)
    p_entry.insert(INSERT,str(data))
#Lưu file đề
def save():
    f = asksaveasfile(mode='w', defaultextension=".txt")
    if f is None:
        return
    text2save = str(ans1_entry.get(1.0, END))
    f.write(text2save)
    f.close()

# Khởi tạo cặp khóa DSA
private_key = dsa.generate_private_key(key_size=1024)
public_key = private_key.public_key()

def Ky():
    message = p_entry.get(1.0, END)
    ans1_entry.delete(1.0,END)
    ans1_entry.insert(INSERT,Kyso(message))
    def read_file(self, file_path):

        if file_path.endswith('.txt'):
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                detected_result = chardet.detect(raw_data)
                detected_encoding = detected_result['encoding']
                file_content = raw_data.decode(detected_encoding)
                return file_content

        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            paragraphs = doc.paragraphs

            formatted_text = []
            for paragraph in paragraphs:
                text = paragraph.text
                formatted_text.append(text)
            return '\n'.join(formatted_text)
        else:
            return "Unsupported file format."
def Xacnhanky():
    # Xác nhận chữ ký
    signature_str = p1_entry.get(1.0, END)
    message = ans2_entry.get(1.0, END)
    p2_entry.delete(1.0,END)
    p2_entry.insert(INSERT,Xacnhankyso(message, signature_str))


def trans():
    text = p_entry.get(1.0, END)
    ans2_entry.delete(1.0, END)
    ans2_entry.insert(INSERT, text)
    ky = ans1_entry.get(1.0, END)
    p1_entry.delete(1.0, END)
    p1_entry.insert(INSERT, ky)

#Các nút bấm
#Mã hóa
getkey_btn = Button(window, text='Lấy L', width=15, height=1, font=('Time New Roman', 10, 'bold'), bg='blue')
getkey_btn.place(x=300, y=35)

upfile1_btn = Button(window, text='File', width=10, height=1, font=('Time New Roman', 10, 'bold'), bg='blue')
upfile1_btn.place(x=500, y=130)

encrypt_btn = Button(window, text='Sinh chữ ký', width=10, height=1, font=('Time New Roman', 10, 'bold'), bg='blue', command = Ky)
encrypt_btn.place(x=150, y=270)

trans_btn = Button(window, text='Chuyển', width=10, height=1, font=('Time New Roman', 10, 'bold'), bg='blue', command = trans)
trans_btn.place(x=500, y=350)

save_btn = Button(window, text='Lưu', width=10, height=1, font=('Time New Roman', 10, 'bold'), bg='blue', command=save)
save_btn.place(x=500, y=400)

#Giải mã

upfile2_btn = Button(window, text='File', width=10, height=1, font=('Time New Roman', 10, 'bold'), bg='blue')
upfile2_btn.place(x=1200, y=130)

decrypt_btn = Button(window, text='Kiểm tra', width=10, height=1, font=('Time New Roman', 10, 'bold'), bg='blue',command=Xacnhanky)
decrypt_btn.place(x=850, y=270)

desave_btn = Button(window, text='Lưu', width=10, height=1, font=('Time New Roman', 10, 'bold'), bg='blue')
desave_btn.place(x=1200, y=370)


window.mainloop()